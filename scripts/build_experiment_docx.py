from pathlib import Path
import shutil

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
ASSETS = RESULTS / "report_assets"
OUT = ROOT / "高性能计算光线追踪渲染器实验报告.docx"

SCREENSHOT_SERIAL = Path(
    "/Users/canghe/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_fbhucv2so65n12_7368/temp/RWTemp/2026-07/"
    "b89452a51d35f5f02eade11c067c81c4.png"
)
SCREENSHOT_OPENMP = Path(
    "/Users/canghe/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_fbhucv2so65n12_7368/temp/RWTemp/2026-07/"
    "75f51280-6bfb-49be-9625-6d25eab7a32a/1.png"
)


def ensure_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    copies = {}
    for src, name in [
        (SCREENSHOT_SERIAL, "terminal_serial.png"),
        (SCREENSHOT_OPENMP, "terminal_openmp.png"),
        (RESULTS / "hpc_output_1.png", "hpc_output_1.png"),
        (RESULTS / "hpc_output_3.png", "hpc_output_3.png"),
        (RESULTS / "showcase_blackhole_preview.png", "showcase_blackhole_preview.png"),
        (RESULTS / "showcase_city_drive_preview.png", "showcase_city_drive_preview.png"),
        (RESULTS / "showcase_snow_gt_preview.png", "showcase_snow_gt_preview.png"),
    ]:
        if src.exists():
            dst = ASSETS / name
            shutil.copyfile(src, dst)
            copies[name] = dst

    for src, name in [
        (RESULTS / "hpc_output_1.ppm", "hpc_output_1.jpg"),
        (RESULTS / "hpc_output_3.ppm", "hpc_output_3.jpg"),
        (RESULTS / "v3_output.ppm", "v3_output.jpg"),
    ]:
        if src.exists():
            dst = ASSETS / name
            im = Image.open(src).convert("RGB")
            im.save(dst, "JPEG", quality=95)
            copies[name] = dst
            png_dst = ASSETS / name.replace(".jpg", ".png")
            im.save(png_dst, "PNG")
            copies[png_dst.name] = png_dst
    return copies


def set_east_asia_font(run_or_style, font_name):
    element = run_or_style._element
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa="8600"):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), width_dxa)
    tbl_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    set_east_asia_font(normal, "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)

    heading1 = styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    set_east_asia_font(heading1, "黑体")
    heading1.font.size = Pt(15)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    set_east_asia_font(heading2, "黑体")
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2.paragraph_format.space_before = Pt(8)
    heading2.paragraph_format.space_after = Pt(4)


def add_header_footer(doc):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "并行编程实验设计文档"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "第 8 组  24281098 胡哲祺 / 24281100 李建宇"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for para in (header_p, footer_p):
        for run in para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            set_east_asia_font(run, "宋体")


def add_center(doc, text, size=12, bold=False, before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    set_east_asia_font(r, "宋体" if not bold else "黑体")
    return p


def add_cover(doc):
    add_center(doc, "并行编程实验设计文档", size=26, bold=True, before=70, after=36)
    cover_items = [
        ("实验项目名称：", "光线追踪渲染器的串行、OpenMP 与 CUDA 并行加速"),
        ("姓        名：", "胡哲祺、李建宇"),
        ("班        级：", "计科2402"),
        ("学        号：", "24281098、24281100"),
        ("小组序号：", "第 8 组"),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(1.05)
        p.paragraph_format.space_after = Pt(12)
        r1 = p.add_run(label)
        r1.font.size = Pt(14)
        set_east_asia_font(r1, "宋体")
        r2 = p.add_run(value)
        r2.font.size = Pt(14)
        set_east_asia_font(r2, "宋体")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.05)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("自我评价：")
    r.font.size = Pt(14)
    set_east_asia_font(r, "宋体")
    for text in [
        "本课程设计围绕游戏画面渲染中的光线追踪计算展开，完成了串行基准、OpenMP 多线程优化和 CUDA GPU 并行优化三个版本。",
        "实验能够在超算平台完成编译运行，结果包含运行截图、渲染输出和性能对比，基本达到课程设计对代码复现、结果展示和性能分析的要求。",
    ]:
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(1.05)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Pt(21)
        p.paragraph_format.space_after = Pt(4)

    add_center(doc, "成绩：___________", size=14, before=24, after=54)
    add_center(doc, "北京交通大学计算机与信息技术学院", size=14, bold=True, before=20, after=10)
    add_center(doc, "2026 年 7 月", size=14, after=0)
    doc.add_page_break()


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.add_run(item)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-12)
        p.paragraph_format.space_after = Pt(3)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    set_cell_border(cell, "C9C9C9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Menlo"
        set_east_asia_font(r, "Menlo")
        r.font.size = Pt(8.5)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(80, 80, 80)
    set_east_asia_font(r, "宋体")


def add_image(doc, path, caption, width=5.9):
    if not path or not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_two_images(doc, left_path, right_path, left_caption, right_caption):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, "FFFFFF")
    for idx, (path, caption) in enumerate([(left_path, left_caption), (right_path, right_caption)]):
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path and Path(path).exists():
            p.add_run().add_picture(str(path), width=Inches(2.9))
        cap_p = table.cell(1, idx).paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap_p.add_run(caption)
        r.font.size = Pt(9)
        r.italic = True
        set_east_asia_font(r, "宋体")


def add_perf_table(doc):
    serial = 39.1013
    openmp = 9.72383
    cuda = 0.126103
    openmp_speedup = serial / openmp
    cuda_speedup = serial / cuda
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    headers = ["版本", "并行方式", "测试参数", "运行时间/s", "相对串行加速比"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "EDEDED")
    rows = [
        ["V1 Serial", "单线程 CPU", "400×300, 64 samples", f"{serial:.4f}", "1.00"],
        ["V2 OpenMP", "8 线程 CPU", "400×300, 64 samples", f"{openmp:.5f}", f"{openmp_speedup:.2f}"],
        ["V3 CUDA", "GPU 并行", "400×300, 64 samples", f"{cuda:.6f}", f"{cuda_speedup:.2f}"],
    ]
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
    add_caption(doc, "表 1  串行、OpenMP 与 CUDA 三版本超算实测性能对比")


def build_doc():
    assets = ensure_assets()
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_cover(doc)

    doc.add_heading("一、实验目的", level=1)
    add_para(
        doc,
        "本实验以光线追踪渲染器为研究对象，围绕游戏画面渲染中常见的像素级独立计算任务，完成串行基准程序、OpenMP CPU 多线程版本和 CUDA GPU 并行版本的设计、实现与性能测试。",
    )
    add_para(
        doc,
        "通过本实验，理解光线追踪中射线生成、相交判断、材质散射、随机采样和图像输出的基本过程，掌握使用高性能计算技术优化计算密集型程序的常见方法，并能够用运行截图、渲染结果和加速比数据评价优化效果。",
    )

    doc.add_heading("二、实验环境", level=1)
    doc.add_heading("2.1 平台环境", level=2)
    add_bullets(
        doc,
        [
            "本地开发环境：macOS，用于代码整理、文档编写、结果图转换和压缩包归档。",
            "超算登录节点：yinxiu.nsccwx.cn，账号 bjtu1/bjtu3，工作目录 data/24281100/1 与 data/24281100/2。",
            "超算 GPU 作业分区：gpu_4090，CUDA 程序通过 Slurm 作业脚本提交运行。",
        ],
    )
    doc.add_heading("2.2 软件环境", level=2)
    add_bullets(
        doc,
        [
            "C++ 编译器：g++，使用 -O2 -std=c++17 编译串行版本。",
            "OpenMP 编译选项：g++ -O2 -std=c++17 -fopenmp。",
            "CUDA 编译环境：module load intel/cuda/12.1，nvcc -O3 -arch=sm_80。",
            "结果格式：程序输出 PPM 图像，报告中将 PPM 转换为 JPG/PNG 以便展示。",
        ],
    )

    doc.add_heading("三、实验内容及程序设计", level=1)
    doc.add_heading("3.1 实验一：串行与 OpenMP 光线追踪性能比较", level=2)
    add_para(
        doc,
        "串行版本作为基准程序，对每个像素按顺序进行多次随机采样，再调用 ray_color 完成光线递归追踪。OpenMP 版本利用像素之间相互独立的特点，对图像行循环进行并行化。",
    )
    add_code(
        doc,
        """
// 串行版本：逐行逐像素计算颜色
for (int j = HEIGHT - 1; j >= 0; j--) {
    for (int i = 0; i < WIDTH; i++) {
        Vec3 color(0, 0, 0);
        for (int s = 0; s < samples; s++) {
            Ray r = camera.get_ray(u, v);
            color += ray_color(r, world, max_depth);
        }
        write_pixel(out, color, samples);
    }
}

// OpenMP 版本：先写入 framebuffer，最后统一输出文件
#pragma omp parallel for schedule(dynamic, 1)
for (int j = 0; j < height; j++) {
    for (int i = 0; i < width; i++) {
        framebuffer[j * width + i] = trace_pixel(i, j);
    }
}
""",
    )
    add_code(
        doc,
        """
g++ -O2 -std=c++17 v1_serial/main.cpp -o v1_serial/v1_serial
g++ -O2 -std=c++17 -fopenmp v2_openmp/main.cpp -o v2_openmp/v2_openmp
./v1_serial/v1_serial --width 400 --height 300 --samples 64
OMP_NUM_THREADS=8 ./v2_openmp/v2_openmp --width 400 --height 300 --samples 64
""",
    )

    doc.add_heading("3.2 实验二：CUDA GPU 光线追踪并行加速", level=2)
    add_para(
        doc,
        "CUDA 版本将一个像素映射为一个 GPU 线程，用二维 grid/block 组织线程。场景球体和材质数据放入显存后传入 kernel，避免在 __constant__ 变量中使用带动态初始化的对象。",
    )
    add_code(
        doc,
        """
__global__ void render_kernel(float* framebuffer, int width, int height,
                              int samples, Camera camera,
                              curandState* rand_states,
                              const Sphere* spheres,
                              const MaterialData* materials) {
    int i = blockIdx.x * blockDim.x + threadIdx.x;
    int j = blockIdx.y * blockDim.y + threadIdx.y;
    if (i >= width || j >= height) return;

    int idx = j * width + i;
    framebuffer[3 * idx + 0] = color.x();
    framebuffer[3 * idx + 1] = color.y();
    framebuffer[3 * idx + 2] = color.z();
}
""",
    )
    add_code(
        doc,
        """
module load intel/cuda/12.1
nvcc -O3 -arch=sm_80 v3_cuda/main.cu -o v3_cuda/v3_cuda
sbatch submit.sh ./v3_cuda/v3_cuda
squeue -u bjtu3
""",
    )

    doc.add_heading("3.3 扩展展示场景设计", level=2)
    add_para(
        doc,
        "为了使课程设计结果更贴近游戏画面渲染方向，项目额外加入了 racing、blackhole、snow_gt 等场景参数。车辆、公路、雪地和黑洞吸积盘等画面均由程序化几何和材质生成，不依赖商业游戏源代码或贴图资源。",
    )

    doc.add_heading("四、实验结果", level=1)
    doc.add_heading("4.1 实验一运行结果", level=2)
    add_code(
        doc,
        """
./v1_serial/v1_serial --width 400 --height 300 --samples 64
Rendering scanline 300/300
Done.
Render time: 39.1013 s

OMP_NUM_THREADS=8 ./v2_openmp/v2_openmp --width 400 --height 300 --samples 64
Threads=8 Render time: 9.72383 s
""",
    )
    add_image(doc, assets.get("terminal_serial.png"), "图 1  串行版本超算编译与运行截图", 5.9)
    add_image(doc, assets.get("terminal_openmp.png"), "图 2  OpenMP 版本运行、结果文件检查与运行时间截图", 5.9)

    doc.add_heading("4.2 实验二运行结果", level=2)
    add_code(
        doc,
        """
CUDA Render time: 0.126103 s
Output: results/v3_output.ppm
Environment: intel/cuda/12.1, gpu_4090 partition
""",
    )
    add_image(doc, assets.get("v3_output.jpg"), "图 3  CUDA 版本输出 v3_output.ppm 转换后的渲染结果", 4.8)
    add_perf_table(doc)

    doc.add_heading("4.3 渲染效果展示", level=2)
    add_para(doc, "下图为从超算结果和程序化展示场景中选取的代表性图片，用于说明程序不仅可以完成性能测试，也能生成具有游戏画面风格的可视化结果。")
    add_two_images(
        doc,
        assets.get("hpc_output_1.jpg"),
        assets.get("hpc_output_3.jpg"),
        "图 4(a)  超算输出 hpc_output_1.ppm",
        "图 4(b)  超算输出 hpc_output_3.ppm",
    )
    add_two_images(
        doc,
        assets.get("showcase_city_drive_preview.png"),
        assets.get("showcase_snow_gt_preview.png"),
        "图 5(a)  城市追车风格展示图",
        "图 5(b)  雪地赛道风格展示图",
    )
    add_image(doc, assets.get("showcase_blackhole_preview.png"), "图 6  黑洞吸积盘风格展示图", 5.7)

    doc.add_heading("五、结果分析与问题说明", level=1)
    doc.add_heading("5.1 OpenMP 加速效果分析", level=2)
    add_para(
        doc,
        "OpenMP 8 线程版本将运行时间由 39.1013 s 降低到 9.72383 s，加速比约为 4.02。由于文件输出、场景初始化和部分随机数处理仍存在串行开销，并且不同像素的反射/折射次数不完全一致，因此没有达到理想的 8 倍加速。",
    )
    add_code(
        doc,
        """
OpenMP Speedup = 39.1013 / 9.72383 ≈ 4.02
OpenMP Efficiency = 4.02 / 8 ≈ 50.3%
""",
    )
    doc.add_heading("5.2 CUDA 加速效果分析", level=2)
    add_para(
        doc,
        "CUDA 版本将每个像素交给独立 GPU 线程处理，在 400×300、64 samples 的测试参数下运行时间为 0.126103 s，相对串行版本加速约 310.07 倍。该结果说明光线追踪中的像素级并行度很高，适合使用 GPU 大规模并行计算。",
    )
    add_code(
        doc,
        """
CUDA Speedup = 39.1013 / 0.126103 ≈ 310.07
""",
    )
    doc.add_heading("5.3 两实验对比总结", level=2)
    add_para(
        doc,
        "CPU 多线程优化实现简单、调试方便，适合作为第一步并行化方案；GPU 优化对内存布局、线程组织和设备运行环境要求更高，但在像素数量较多、采样次数较高时能够获得更突出的加速效果。实验过程中曾出现登录节点 CUDA runtime 与 driver 不匹配的问题，最终通过 Slurm 提交至 gpu_4090 分区解决。",
    )

    doc.add_heading("六、实验体会", level=1)
    add_para(
        doc,
        "通过本次课程设计，我们从一个串行光线追踪程序出发，逐步完成 OpenMP 和 CUDA 两种并行优化。实验体会最深的是：高性能计算不仅是添加并行语句，还需要考虑任务划分、内存访问、随机数状态、输出同步和运行平台差异。",
    )
    add_para(
        doc,
        "本项目最终能够在超算平台生成渲染图片并取得明显加速，说明光线追踪渲染是展示高性能计算价值的合适案例。后续若继续完善，可加入 OBJ 模型加载、BVH 加速结构和更复杂材质，从而进一步接近真实游戏渲染管线。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)

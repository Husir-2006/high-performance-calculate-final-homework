from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "高性能计算光线追踪渲染器答辩稿.docx"


SLIDES = [
    (
        "第 1 页：封面",
        "各位老师好，我们是第 8 组，组员是 24281098 胡哲祺和 24281100 李建宇。"
        "我们本次课程设计的题目是光线追踪渲染器的并行优化。"
        "项目主要围绕一个可运行的光线追踪程序展开，分别实现了串行版本、OpenMP 多线程版本和 CUDA GPU 版本。"
        "从最终测试结果看，在 400 乘 300 分辨率、64 次采样的参数下，OpenMP 8 线程相对串行加速约 4.02 倍，CUDA 版本相对串行加速约 310.07 倍。"
        "接下来我会按背景、实现方法、运行结果和性能分析几个部分介绍我们的工作。",
    ),
    (
        "第 2 页：课程设计任务与提交要求",
        "这一页对应课程设计要求。老师要求我们自选题目，用高性能计算技术优化已有代码、算法或程序，并提交源代码、实验报告和答辩 PPT。"
        "所以我们选择了光线追踪渲染器作为优化对象。它本身是一个计算量比较大的图形程序，而且每个像素之间基本独立，非常适合用并行计算加速。"
        "本项目保留了原始串行代码，同时给出了 OpenMP 和 CUDA 两个优化版本，并整理了复现说明，方便助教老师在超算平台上重新编译和运行。",
    ),
    (
        "第 3 页：为什么选择光线追踪",
        "选择光线追踪主要有两个原因。第一，光线追踪能表现反射、折射、阴影和景深等效果，和游戏画面、影视渲染都有关系，展示效果比较直观。"
        "第二，它的计算模式很适合并行化。每个像素都需要发射光线、和场景求交、根据材质计算散射，并且通常要多次采样。"
        "这些像素之间没有强依赖关系，因此可以把像素任务分配给多个 CPU 线程，或者进一步映射到 GPU 上的大量线程。"
        "所以这个题目既能体现高性能计算思想，也能用图片结果直接验证程序是否正确。",
    ),
    (
        "第 4 页：光线追踪核心流程",
        "这一页展示的是光线追踪的基本流程。首先，相机会根据像素位置生成一条 primary ray，也就是从相机出发的光线。"
        "然后程序会在场景中查找这条光线和物体的最近交点，并得到交点法线。接着根据材质类型决定光线如何散射，例如漫反射、金属反射或者玻璃折射。"
        "为了降低锯齿和噪声，每个像素会做多次随机采样，最后把这些采样颜色平均后写入图像。"
        "这个流程中，每个像素都要重复大量类似计算，因此优化空间主要集中在像素循环的并行化上。",
    ),
    (
        "第 5 页：技术路线",
        "我们的技术路线是从 V1 到 V3 逐步优化。V1 是串行版本，按顺序逐像素渲染，作为正确性和性能基准。"
        "V2 是 OpenMP 版本，在 CPU 上对像素循环进行多线程并行。V3 是 CUDA 版本，把每个像素映射到一个 GPU thread。"
        "三种版本的核心渲染算法保持一致，主要变化是任务调度方式。这样做的好处是可以比较清楚地分析，性能提升到底来自 CPU 多线程，还是来自 GPU 的大规模并发。",
    ),
    (
        "第 6 页：代码结构",
        "程序结构主要围绕几个核心模块组织。vec3 负责三维向量运算，ray 表示光线，sphere 负责球体求交，material 负责材质散射，camera 负责相机模型，main 文件负责场景构建、主循环和输出。"
        "串行版本和 OpenMP 版本使用 C++ 头文件，CUDA 版本对应使用 cuh 文件和 main.cu。"
        "这种结构让三个版本可以共享同一套光线追踪思想，同时又能针对 CPU 和 GPU 分别改写实现细节。",
    ),
    (
        "第 7 页：串行版本",
        "串行版本是整个项目的基准程序。它按照图像行和列逐个像素计算，每个像素执行指定次数的随机采样，然后调用 ray_color 函数递归追踪光线。"
        "串行版本的优点是逻辑清晰，便于调试和验证图像是否正确；缺点是只能使用一个 CPU 核心。"
        "在我们的超算实测中，400 乘 300 分辨率、64 次采样时，串行版本运行时间是 39.1013 秒。后续所有加速比都以这个时间作为基准。",
    ),
    (
        "第 8 页：OpenMP 版本",
        "OpenMP 版本利用了像素之间相互独立的特点，对图像行循环加 parallel for。为了避免多个线程同时写同一个文件，我们先把计算结果写入 framebuffer，最后再按顺序输出 PPM 文件。"
        "这里还使用了 schedule dynamic，这是因为不同像素的反射和折射次数可能不同，计算量不完全一致。动态调度可以让先完成任务的线程继续领取新任务，缓解负载不均。"
        "在 8 线程下，OpenMP 版本运行时间为 9.72383 秒，相对串行加速约 4.02 倍。",
    ),
    (
        "第 9 页：CUDA 版本",
        "CUDA 版本进一步把像素计算映射到 GPU 线程。我们使用 16 乘 16 的线程块组织二维网格，每个 CUDA thread 负责一个像素的采样和颜色计算。"
        "为了适应 GPU，递归形式的 ray_color 被改成迭代形式，减少 GPU 栈压力。"
        "球体和材质数据通过 cudaMalloc 分配到显存中，再作为 kernel 参数传入。"
        "最终在超算 gpu_4090 分区上运行，CUDA 版本运行时间为 0.126103 秒，相对串行加速约 310.07 倍。",
    ),
    (
        "第 10 页：CUDA 内存问题处理",
        "开发 CUDA 版本时遇到过一个比较关键的问题。早期代码尝试把 Sphere 和 MaterialData 数组放到 __constant__ 变量中，但这些结构体间接包含 Vec3 构造逻辑，nvcc 会报动态初始化不支持。"
        "后来我们改成在 host 端先构建场景，再用 cudaMalloc 分配显存，用 cudaMemcpy 把球体和材质数组复制到 GPU。"
        "kernel 启动时显式传入 spheres、materials 和 num_spheres。这样既解决了编译问题，也让 CUDA 版本的数据传递更稳定。",
    ),
    (
        "第 11 页：超算平台运行流程",
        "这一页展示的是超算平台上的运行流程。首先进入项目目录，然后加载 CUDA 模块 intel/cuda/12.1，再用 nvcc 编译 main.cu。"
        "需要注意的是，登录节点不能直接作为 GPU 运行环境。我们之前在登录节点直接运行 CUDA 程序时遇到过 driver version is insufficient 的错误。"
        "这个错误不是程序逻辑问题，而是登录节点没有匹配的 GPU 驱动环境。正确做法是通过 sbatch 把任务提交到 gpu_4090 分区，在 GPU 计算节点上运行。",
    ),
    (
        "第 12 页：运行结果",
        "这一页是三种版本的运行时间对比。串行版本是 39.1013 秒，OpenMP 8 线程版本是 9.72383 秒，CUDA 版本是 0.126103 秒。"
        "从图中可以看到，OpenMP 已经能显著降低运行时间，而 CUDA 的运行时间进一步下降到 1 秒以内。"
        "按照串行版本作为基准，OpenMP 加速约 4.02 倍，CUDA 加速约 310.07 倍。这个结果说明光线追踪这种像素级任务非常适合 GPU 并行。",
    ),
    (
        "第 13 页：性能分析",
        "从性能分析角度看，OpenMP 的收益来自 CPU 多核并行，但它没有达到理想的 8 倍线性加速。原因包括文件输出和部分初始化仍然是串行部分，不同像素计算量不同会带来负载不均，随机数生成和动态调度也会有额外开销。"
        "CUDA 版本加速更明显，是因为 GPU 可以同时调度大量线程，每个像素独立计算，能够充分利用 GPU 的并发能力。"
        "当然 CUDA 也会受到分支发散、随机数状态访问和显存传输等因素影响，但总体上更适合这种大规模像素计算。",
    ),
    (
        "第 14 页：结果图片",
        "这一页展示的是 PPM 输出和转换后的结果图。PPM 格式结构简单，非常适合在超算环境中直接生成和检查。"
        "为了放入报告和 PPT，我们把关键 PPM 文件转换成 JPG 和 PNG。"
        "这些图像一方面用于验证串行、OpenMP 和 CUDA 版本都能正常输出，另一方面也展示了程序化场景的视觉效果。",
    ),
    (
        "第 15 页：雪地赛车展示",
        "这一页是雪地赛车场景的参考展示。我们的方向是游戏画面渲染，所以除了基础小球场景，还加入了更接近游戏截图风格的展示场景。"
        "这个场景强调高速运动、雪地反射、车身高光和背景运动感。"
        "当前程序主要用程序化几何和像素着色近似这些效果，还没有使用商业游戏素材，因此可以避免版权问题。",
    ),
    (
        "第 16 页：城市追车展示",
        "这一页是城市追车风格展示。这里主要想表现城市道路、秋日光照、反射和景深氛围。"
        "需要说明的是，PPT 中的高清画面是答辩参考图，用来说明我们希望靠近的游戏渲染方向；项目中的程序输出是由球体、材质和像素着色过程生成的。"
        "这样既能让展示更直观，也能说明后续如果加入 OBJ 模型和 BVH 加速，画面真实感还可以继续提升。",
    ),
    (
        "第 17 页：黑洞场景展示",
        "黑洞场景主要用于展示程序化像素着色能力。它不依赖复杂三维模型，而是通过像素函数、噪声、环形结构和高亮边缘来模拟吸积盘和黑洞轮廓。"
        "这个场景比较适合 CUDA 版本，因为每个像素都可以独立计算颜色。"
        "从展示角度看，它也能让答辩材料更突出，不只是普通的基础测试图。",
    ),
    (
        "第 18 页：复现方法",
        "这一页列出了助教复现时可以使用的核心命令。CPU 部分用 g++ 编译串行和 OpenMP 版本，CUDA 部分加载 intel/cuda/12.1 后用 nvcc 编译。"
        "OpenMP 运行时可以通过 OMP_NUM_THREADS 指定线程数。CUDA 运行时建议使用 submit.sh 提交到 GPU 分区。"
        "提交包里还包含 SUBMISSION_README、EXPERIMENT_GUIDE 和 SUPERCOMPUTER_GUIDE，分别对应最短复现说明、完整实验步骤和超算环境说明。",
    ),
    (
        "第 19 页：不足与后续改进",
        "目前项目还可以继续改进。第一，可以加入 OBJ 模型加载，把车辆从程序化球体近似过渡到真实三角网格。"
        "第二，可以加入 BVH 加速结构，减少大量物体场景下的求交次数。"
        "第三，可以继续测试不同分辨率、采样数和线程数下的扩展性。"
        "这些方向都和实际图形渲染系统比较接近，也是后续可以继续完善的部分。",
    ),
    (
        "第 20 页：总结",
        "最后总结一下，本项目完成了串行、OpenMP 和 CUDA 三个版本的光线追踪渲染器，并在超算平台完成了三版本实测。"
        "OpenMP 8 线程相对串行加速约 4.02 倍，CUDA 相对串行加速约 310.07 倍。"
        "实验说明光线追踪具有明显的像素级并行特征，适合使用高性能计算技术进行优化。"
        "我们也整理了源代码、实验报告、答辩 PPT PDF 和复现说明，满足课程设计提交要求。我的汇报到这里结束，谢谢老师和同学。",
    ),
]


def set_east_asia_font(run_or_style, font_name):
    element = run_or_style._element
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    set_east_asia_font(normal, "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(7)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    set_east_asia_font(h1, "黑体")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(31, 77, 120)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("高性能计算课程设计答辩稿")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 77, 120)
    set_east_asia_font(run, "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("光线追踪渲染器的串行、OpenMP 与 CUDA 并行加速")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(80, 80, 80)
    set_east_asia_font(run, "宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第 8 组：24281098 胡哲祺，24281100 李建宇")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    set_east_asia_font(run, "宋体")


def add_note(doc):
    p = doc.add_paragraph()
    run = p.add_run(
        "使用建议：录制视频时按 PPT 页码逐页朗读。可根据实际语速删减个别句子，"
        "但建议保留三组核心数据：串行 39.1013 s、OpenMP 9.72383 s、CUDA 0.126103 s。"
    )
    run.bold = True
    run.font.color.rgb = RGBColor(120, 65, 0)
    set_east_asia_font(run, "宋体")
    p.paragraph_format.space_after = Pt(10)


def build_doc():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_note(doc)

    for title, script in SLIDES:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(script)
        p.paragraph_format.first_line_indent = Pt(22)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
